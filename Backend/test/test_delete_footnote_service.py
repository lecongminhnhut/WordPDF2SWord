import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.services.delete_footnote_service import DeleteFootnoteService


class DeleteFootnoteServiceTest(unittest.TestCase):
    def setUp(self):
        self.service = DeleteFootnoteService()
        self.temp_dir = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.temp_dir.cleanup()

    def create_document(self, paragraphs):
        path = Path(self.temp_dir.name) / "test.docx"
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    @staticmethod
    def read_paragraphs(path):
        return [paragraph.text for paragraph in Document(path).paragraphs]

    def test_removes_vietnamese_notes_without_deleting_next_chapter(self):
        path = self.create_document([
            "Nội dung cuối chương.",
            "————–",
            "Chú thích:",
            "(1) Chú thích thứ nhất.",
            "(2) Chú thích thứ hai.",
            "",
            "BÚP SEN XANH BÚP SEN XANH",
            "Chương II",
            "Nội dung chương tiếp theo.",
        ])

        result = self.service.remove_footnotes_precisely(path)

        self.assertEqual(result["blocks_removed"], 1)
        self.assertEqual(self.read_paragraphs(path), [
            "Nội dung cuối chương.",
            "",
            "BÚP SEN XANH BÚP SEN XANH",
            "Chương II",
            "Nội dung chương tiếp theo.",
        ])

    def test_supports_common_endnote_markers(self):
        path = self.create_document([
            "Endnotes",
            "[1] First note.",
            "2. Second note.",
            "† Symbol note.",
            "Appendix A",
        ])

        result = self.service.remove_footnotes_precisely(path)

        self.assertEqual(result["blocks_removed"], 1)
        self.assertEqual(self.read_paragraphs(path), ["Appendix A"])

    def test_keeps_numbered_lists_without_note_heading(self):
        original = [
            "Các bước thực hiện:",
            "(1) Chuẩn bị tài liệu.",
            "(2) Kiểm tra tài liệu.",
            "Chương tiếp theo.",
        ]
        path = self.create_document(original)

        result = self.service.remove_footnotes_precisely(path)

        self.assertEqual(result["blocks_removed"], 0)
        self.assertEqual(self.read_paragraphs(path), original)

    def test_keeps_note_heading_without_numbered_items(self):
        original = [
            "Ghi chú",
            "Đây là nội dung chính không sử dụng danh sách chú thích.",
            "Chương tiếp theo.",
        ]
        path = self.create_document(original)

        result = self.service.remove_footnotes_precisely(path)

        self.assertEqual(result["blocks_removed"], 0)
        self.assertEqual(self.read_paragraphs(path), original)

    def test_removes_multiple_note_blocks_independently(self):
        path = self.create_document([
            "Phần một.",
            "Chú thích cuối chương",
            "1) Chú thích phần một.",
            "Phần hai.",
            "Cước chú:",
            "* Chú thích phần hai.",
            "Phần ba.",
        ])

        result = self.service.remove_footnotes_precisely(path)

        self.assertEqual(result["blocks_removed"], 2)
        self.assertEqual(self.read_paragraphs(path), [
            "Phần một.",
            "Phần hai.",
            "Phần ba.",
        ])


if __name__ == "__main__":
    unittest.main()
